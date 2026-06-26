from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# --- Title Page ---
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONSOLATRIX COLLEGE OF TOLEDO CITY, INC.')
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COLLEGE OF COMPUTER STUDIES')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Software Requirements Specifications')
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('for')
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Lokal: A Web Application for Online Hiring and Applicant Management for CJTECH Computer Trading')
run.bold = True
run.font.size = Pt(16)
run.font.italic = True

doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '  1.1 Purpose',
    '  1.2 Scope',
    '  1.3 Definitions, Acronyms and Abbreviations',
    '  1.4 References',
    '2. Overall Description',
    '  2.1 Product Perspective',
    '  2.2 User Characteristics',
    '  2.3 Constraints',
    '  2.4 Assumptions and Dependencies',
    '3. Specific Requirements',
    '  3.1 External Interface Requirements',
    '  3.2 Functional Requirements',
    '  3.3 Non-Functional Requirements',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# --- 1. Introduction ---
doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph(
    'The purpose of this Software Requirements Specification (SRS) is to provide a detailed overview of Lokal, '
    'a web application built for online hiring and applicant management at CJTECH Computer Trading. '
    'This document lays out the functional and non-functional requirements that guided the development, '
    'testing, and deployment of the system. It serves as a reference for developers, testers, project advisers, '
    'and stakeholders so that everyone is on the same page regarding what the system does, why it exists, '
    'and how it is expected to behave.'
)

doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph(
    'Lokal is a web-based application that streamlines and modernizes the way CJTECH Computer Trading handles '
    'recruitment. Job seekers can browse open positions and submit applications online, while administrators '
    'manage job postings, review applicants, and track recruitment activity through a centralized dashboard.'
)
doc.add_paragraph(
    'The platform replaces the old manual hiring process with something faster, more organized, and easier to '
    'access for both sides. On top of that, it uses Artificial Intelligence (AI) through the Google Gemini API '
    'to help recruiters figure out which applicants fit which jobs best, based on the skills listed in their profiles '
    'compared to the job requirements.'
)

p = doc.add_paragraph('Key features of the platform include:')
p.paragraph_format.space_after = Pt(2)

features = [
    'Job Browsing and Search — Applicants can search available vacancies by title or description.',
    'Online Application Submission — Applicants submit applications electronically with data pulled straight from their profiles.',
    'Profile Management — Users can update their name, contact info, location, skills, bio, and resume link at any time.',
    'AI-Based Skill Matching — The system compares applicant skills against job requirements using Google Gemini AI, and falls back to a basic keyword matcher if the AI is unavailable.',
    'Re-apply for Rejected Applications — Applicants whose applications were rejected can re-apply with a single click.',
    'Hire / Reject Decisions — Administrators can mark applicants as hired or rejected directly from the applicant detail view.',
    'Notifications — Applicants receive updates about application status and hiring decisions.',
    'Job Posting Management — Administrators can create, edit, and delete job openings with an icon picker for visual flair.',
    'Applicant Ranking — Applicants are ranked by their AI match score so admins see the best candidates first.',
    'Application Tracking — Applicants can see the status of all their submissions in one place.',
]

for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('1.3 Definitions, Acronyms and Abbreviations', level=2)

table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Term', 'Definition']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

terms = [
    ('Lokal', 'The web-based hiring and applicant management system developed for CJTECH Computer Trading.'),
    ('User', 'Any person using the system, including applicants and administrators.'),
    ('Applicant', 'An individual seeking employment who applies for job vacancies through the system.'),
    ('Admin', 'A user with administrative privileges to manage jobs, applicants, and system records.'),
    ('Job Posting', 'An advertisement containing the position title, description, requirements, salary, and location.'),
    ('Application', 'A submission made by an applicant for a specific job vacancy.'),
    ('AI Skill Match', 'A feature that evaluates applicant skills against job requirements using the Google Gemini API.'),
    ('Notification', 'A system-generated message informing users about application updates and hiring decisions.'),
]

for idx, (term, definition) in enumerate(terms):
    row = table.rows[idx + 1]
    row.cells[0].text = term
    row.cells[1].text = definition
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(11)

doc.add_heading('1.4 References', level=2)
refs = [
    'IEEE Standard for Software Requirements Specifications — IEEE Std 830-1998, Institute of Electrical and Electronics Engineers.',
    'React Documentation (v18) — https://react.dev, Meta Open Source.',
    'Vite Documentation — https://vitejs.dev, Evan You.',
    'Supabase Documentation — https://supabase.com/docs, Supabase Inc.',
    'Google Gemini API Documentation — https://ai.google.dev, Google.',
    'PostgreSQL Documentation — https://www.postgresql.org/docs, PostgreSQL Global Development Group.',
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# --- 2. Overall Description ---
doc.add_heading('2. Overall Description', level=1)

doc.add_heading('2.1 Product Perspective', level=2)
doc.add_paragraph(
    'Lokal is a standalone web application built with React and Vite on the frontend, and Supabase (which provides '
    'PostgreSQL, authentication, and storage) on the backend. The AI matching feature uses the Google Gemini API '
    'with a built-in fallback to basic substring matching so the app still works even if the API key expires or '
    'the network is down.'
)
doc.add_paragraph(
    'The system does not depend on any external recruitment platforms. It is self-contained: user accounts, job data, '
    'applications, and notifications all live inside the Supabase PostgreSQL database. The design is modular enough '
    'that future versions could hook into third-party email services or external job boards without rewriting the core.'
)

p = doc.add_paragraph('Modular Decomposition')
p.runs[0].bold = True
p.runs[0].font.size = Pt(12)

modules = [
    ('Module 1: User Management', [
        '1.1 User Registration (Signup)',
        '1.2 User Authentication (Login / Logout)',
        '1.3 Profile Management (Edit name, phone, location, skills, bio, resume link, avatar)',
    ]),
    ('Module 2: Job Management', [
        '2.1 Browse and Search Jobs',
        '2.2 Create / Edit / Delete Job Postings (Admin)',
        '2.3 View Job Details',
    ]),
    ('Module 3: Application Management', [
        '3.1 Submit Application (data sourced from profile)',
        '3.2 Track Application Status',
        '3.3 Re-apply for Rejected Applications',
    ]),
    ('Module 4: AI-Based Applicant Matching', [
        '4.1 AI Skill Analysis via Google Gemini API',
        '4.2 Applicant Ranking by Match Score',
        '4.3 Basic Keyword Matching Fallback',
    ]),
    ('Module 5: Admin Dashboard', [
        '5.1 Dashboard Overview (job, applicant, and application counts)',
        '5.2 Manage Applicants (view ranked list, hire / reject)',
        '5.3 User Management (view, approve, suspend, delete users)',
        '5.4 Notification Management',
    ]),
]

for mod_name, transactions in modules:
    p = doc.add_paragraph()
    run = p.add_run(mod_name)
    run.bold = True
    run.font.size = Pt(12)
    for t in transactions:
        doc.add_paragraph(t, style='List Bullet')

doc.add_heading('2.2 User Characteristics', level=2)
doc.add_paragraph('The system serves two main types of users:')

p = doc.add_paragraph()
run = p.add_run('1. Applicant')
run.bold = True
doc.add_paragraph('The primary user looking for work. Applicants can:', style='List Bullet')
doc.add_paragraph('Create and manage their own profile (name, phone, location, skills, bio, resume).', style='List Bullet')
doc.add_paragraph('Upload a profile avatar.', style='List Bullet')
doc.add_paragraph('Browse and search job openings.', style='List Bullet')
doc.add_paragraph('Submit applications to jobs they are interested in.', style='List Bullet')
doc.add_paragraph('Track the status of all their submitted applications.', style='List Bullet')
doc.add_paragraph('Re-apply to jobs where their application was rejected.', style='List Bullet')
doc.add_paragraph('Receive notifications about application updates and hiring decisions.', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('2. Administrator')
run.bold = True
doc.add_paragraph('The recruitment staff or system administrator. Admins can:', style='List Bullet')
doc.add_paragraph('Create, edit, and delete job postings with an icon picker.', style='List Bullet')
doc.add_paragraph('View a ranked list of applicants for any job, ordered by AI match score.', style='List Bullet')
doc.add_paragraph('See a detailed comparison of basic keyword matching vs. AI-powered matching.', style='List Bullet')
doc.add_paragraph('Mark applicants as hired or rejected.', style='List Bullet')
doc.add_paragraph('Manage user accounts (view, approve, suspend, delete).', style='List Bullet')
doc.add_paragraph('View dashboard statistics (total jobs, applicants, applications).', style='List Bullet')

doc.add_heading('2.3 Constraints', level=2)

constraints = [
    ('Regulatory Policies', 'Must comply with the Philippine Data Privacy Act of 2012 (RA 10173) for protecting applicant personal information.'),
    ('Hardware Limitations', 'Runs on standard web infrastructure. No specialized hardware is needed, though server load increases with concurrent users.'),
    ('Browser Support', 'The system targets modern browsers (Chrome, Firefox, Edge, Safari). Older browsers may not support all CSS and JavaScript features used.'),
    ('AI API Dependency', 'The AI matching feature depends on the Google Gemini API. If the API key expires or the service is unreachable, the system falls back to basic keyword matching.'),
    ('Internet Requirement', 'A stable internet connection is required for both applicants and administrators to access the platform.'),
    ('Security', 'Passwords are hashed by Supabase Auth. Applicant data is protected through Row-Level Security policies on the database.'),
]

for title, desc in constraints:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.4 Assumptions and Dependencies', level=2)

assumptions = [
    'Users will access the platform through modern web browsers with JavaScript enabled.',
    'Applicants will provide accurate and complete information in their profiles for the AI matching to work effectively.',
    'Administrators will provide accurate job descriptions and requirements to ensure meaningful match scores.',
    'The system assumes a stable internet connection is available for both applicants and administrators.',
    'The Google Gemini API will remain available under the free-tier quota during the project demonstration period.',
    'The Supabase free tier will provide sufficient storage and request capacity for the scale of this capstone project.',
]

for a in assumptions:
    doc.add_paragraph(a, style='List Bullet')

doc.add_page_break()

# --- 3. Specific Requirements ---
doc.add_heading('3. Specific Requirements', level=1)

doc.add_heading('3.1 External Interface Requirements', level=2)

doc.add_heading('3.1.1 Hardware Interfaces', level=3)
doc.add_paragraph(
    'Lokal is designed to run on standard web infrastructure with minimal hardware demands. '
    'The system is hosted on Supabase\'s cloud servers, so no on-premise hardware is required. '
    'End users only need a device with a modern web browser and an internet connection.'
)
doc.add_paragraph('Supported devices:', style='List Bullet')
doc.add_paragraph('Desktop and laptop computers with screens 1024px or wider for the full dashboard experience.', style='List Bullet')
doc.add_paragraph('Tablets and smartphones — the interface uses responsive CSS so it remains usable on smaller screens.', style='List Bullet')

doc.add_heading('3.1.2 Software Interfaces', level=3)
doc.add_paragraph('The system relies on the following software components:')

interfaces = [
    ('Frontend Framework', 'React 18 with Vite as the build tool. Components are written as functional components using React hooks.'),
    ('Backend / Database', 'Supabase provides PostgreSQL database, user authentication, file storage (for avatars), and Row-Level Security.'),
    ('AI Integration', 'Google Gemini API (gemini-2.0-flash-lite model) for skill matching. The application calls the API through a utility module that wraps the fetch call and handles errors gracefully.'),
    ('Styling', 'Inline styles with a centralized theme file (uiStyles.js) plus a styles.css file for global styles, animations, and utility classes.'),
]

for title, desc in interfaces:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.1.3 Communications Interfaces', level=3)
doc.add_paragraph(
    'All communication between the frontend and Supabase happens over HTTPS using the Supabase JavaScript client library. '
    'The AI matching utility module communicates with the Google Gemini REST API over HTTPS. '
    'No real-time WebSocket connections are currently used, though Supabase has built-in real-time capabilities '
    'that could be enabled in a future version.'
)

doc.add_heading('3.2 Functional Requirements', level=2)

# Module 1
doc.add_heading('Module 1 — User Management', level=3)

doc.add_heading('1.1 User Registration (Signup)', level=4)
doc.add_paragraph('Description: A guest user provides their full name, email address, and password to create an account. An optional phone number can also be provided.')

p = doc.add_paragraph()
run = p.add_run('Preconditions: ')
run.bold = True
p.add_run('The user must not already have an account with the same email. All required fields must be filled.')

p = doc.add_paragraph()
run = p.add_run('Postconditions: ')
run.bold = True
p.add_run('The user is registered but must sign in manually (the system signs them out immediately after signup so they can verify their credentials). A welcome notification is created.')

p = doc.add_paragraph('Main Success Scenario:')
steps = [
    'The guest navigates to the signup page.',
    'The user fills in full name, email, password, and optionally a phone number.',
    'The system validates the inputs.',
    'The user\'s account is created in Supabase Auth.',
    'A corresponding row is inserted into the profiles table.',
    'A welcome notification is inserted.',
    'The user is signed out automatically and shown a success page with a checkmark animation.',
    'The user clicks "Go to Sign In" to proceed to the login page.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('1.2 User Authentication (Login / Logout)', level=4)
doc.add_paragraph('Description: A registered user signs in with their email and password. After successful authentication, the system loads the user\'s profile and redirects them to the appropriate dashboard.')

p = doc.add_paragraph()
run = p.add_run('Preconditions: ')
run.bold = True
p.add_run('The user must have a registered account.')

p = doc.add_paragraph()
run = p.add_run('Postconditions: ')
run.bold = True
p.add_run('The user is authenticated and redirected to their dashboard (admin → /admin, applicant → /find-jobs).')

p = doc.add_paragraph('Main Success Scenario:')
login_steps = [
    'The user enters their email and password on the login page.',
    'The system authenticates via Supabase Auth.',
    'The role is fetched via the get_my_role database function.',
    'A checkmark animation plays and a "Welcome back, [Name]!" message appears.',
    'After a brief pause, the user is redirected to the appropriate dashboard.',
]
for i, step in enumerate(login_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('1.3 Profile Management', level=4)
doc.add_paragraph('Description: Users can view and edit their profile information, including full name, phone number, location, skills, bio, resume link, and avatar.')

p = doc.add_paragraph('Profile fields:')
profile_fields = [
    'Full Name — displayed on the navbar, dashboard sidebar, and applicant cards.',
    'Email — read-only, pulled from Supabase Auth.',
    'Phone Number — optional contact number.',
    'Location — e.g., "Toledo City, Cebu".',
    'Skills — comma-separated list used for AI matching (e.g., "Customer Service, Sales, Communication").',
    'Bio — short description about the applicant.',
    'Resume — a Google Drive link (the app does not store uploaded files due to Supabase free-tier storage limits).',
    'Avatar — profile picture uploaded to Supabase Storage, with cache-busting for immediate updates.',
]
for f in profile_fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# Module 2
doc.add_heading('Module 2 — Job Management', level=3)

doc.add_heading('2.1 Browse and Search Jobs', level=4)
doc.add_paragraph('Description: Applicants can browse all available job postings on the Find Jobs page. A search bar filters jobs by title or description in real time.')

p = doc.add_paragraph()
run = p.add_run('Preconditions: ')
run.bold = True
p.add_run('The applicant must be logged in.')

p = doc.add_paragraph()
run = p.add_run('Postconditions: ')
run.bold = True
p.add_run('The applicant sees a list of matching jobs and can click any job to view details and apply.')

doc.add_heading('2.2 Create, Edit, Delete Job Postings (Admin)', level=4)
doc.add_paragraph('Description: Administrators can create new job postings, edit existing ones, or delete them. A modal form opens when creating or editing, and the job list is displayed in a table.')

p = doc.add_paragraph('Job fields:')
job_fields = [
    'Title — the position name (e.g., "Sales Assistant").',
    'Icon — selected from a grid of 16 business-relevant emojis (e.g., 🏪 🛍️ 💻 🖥️ 🔧).',
    'Company — defaults to "CJTECH Computer Trading".',
    'Location — defaults to "Sangi, Toledo City".',
    'Salary — optional compensation information.',
    'Description — detailed job description.',
    'Requirements — comma-separated list of required skills (e.g., "Customer Service, Basic Accounting, POS System").',
]
for f in job_fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.3 View Job Details', level=4)
doc.add_paragraph('Description: Clicking a job card opens the Apply page showing the full job description, requirements, and the applicant\'s current profile summary side by side.')

# Module 3
doc.add_heading('Module 3 — Application Management', level=3)

doc.add_heading('3.1 Submit Application', level=4)
doc.add_paragraph('Description: Applicants apply for a job by clicking the "Apply Now" button. The application pulls data directly from the applicant\'s saved profile — no separate form to fill out.')

p = doc.add_paragraph()
run = p.add_run('Preconditions: ')
run.bold = True
p.add_run('The applicant must have a complete profile (name, location, and skills filled in).')

p = doc.add_paragraph()
run = p.add_run('Postconditions: ')
run.bold = True
p.add_run('A new application record is created with status "pending". A notification is sent to the applicant.')

doc.add_heading('3.2 Track Application Status', level=4)
doc.add_paragraph('Description: The My Applications page lists all of the applicant\'s submissions with their current status — pending, reviewed, interviewed, hired, or rejected.')

p = doc.add_paragraph('Status definitions:')
statuses = [
    'Pending — the application has been submitted and is awaiting review.',
    'Reviewed — the admin has looked at the application.',
    'Interviewed — the applicant has been scheduled for or has completed an interview.',
    'Hired — the applicant has been accepted for the position.',
    'Rejected — the applicant was not selected.',
]
for s in statuses:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('3.3 Re-apply for Rejected Applications', level=4)
doc.add_paragraph('Description: If an application was rejected, the applicant can click "Apply Again" to reset its status back to "pending" so the admin can review it once more.')

# Module 4
doc.add_heading('Module 4 — AI-Based Applicant Matching', level=3)

doc.add_heading('4.1 AI Skill Analysis', level=4)
doc.add_paragraph(
    'When an admin views applicants for a specific job, the system processes each applicant one at a time with a '
    'visible progress indicator showing "Analyzing (2/4) — Name". For each applicant, the system calls the Google Gemini API '
    'with a prompt that includes the job title, requirements, and the applicant\'s skills.'
)
doc.add_paragraph(
    'The API returns a JSON object containing:'
)
api_fields = [
    'score — an overall match percentage (0–100).',
    'matched — array of skill keywords that matched.',
    'missing — array of required skills the applicant is missing.',
    'explanation — a short plain-text explanation of why the score was given.',
]
for f in api_fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('4.2 Applicant Ranking', level=4)
doc.add_paragraph(
    'After all applicants are analyzed, they are sorted by their AI match score in descending order so the best '
    'candidates appear at the top. The admin sees each applicant\'s score in a colored circle (green for 80%+, '
    'amber for 50–79%, red for below 50%) along with a progress bar.'
)

doc.add_heading('4.3 Basic Keyword Matching Fallback', level=4)
doc.add_paragraph(
    'If the Gemini API call fails (network error, expired key, rate limit), the system silently falls back to '
    'a local keyword-matching algorithm. This algorithm splits skills and requirements into words, counts how many '
    'requirement words appear in the applicant\'s skills, and returns a percentage. The admin never sees an error — '
    'the system just shows the basic match without the AI label.'
)

doc.add_page_break()

# Module 5
doc.add_heading('Module 5 — Admin Dashboard', level=3)

doc.add_heading('5.1 Dashboard Overview', level=4)
doc.add_paragraph('Description: The admin home page shows three stat cards with counts of total jobs, total applicants, and total applications. These numbers are fetched directly from the database using count queries.')

doc.add_heading('5.2 Manage Applicants (Hire / Reject)', level=4)
doc.add_paragraph(
    'Description: From the job detail view, admins can click any applicant to open a detail panel showing their '
    'profile information, AI match explanation, a comparison of basic vs. AI scores, and action buttons to '
    'either hire or reject the applicant. Hiring or rejecting updates the application status in the database '
    'and triggers a notification to the applicant.'
)

doc.add_heading('5.3 User Management', level=4)
doc.add_paragraph(
    'Description: The Users page (accessible from the admin sidebar) lists every registered user with their name, '
    'email, role, status, skills, and resume link. Admins can approve pending users, suspend active users, or '
    'delete accounts entirely.'
)

doc.add_heading('5.4 Notifications', level=4)
doc.add_paragraph(
    'Description: The system automatically creates notification records when key events happen — welcome message '
    'on signup, application submission, status changes, and hiring decisions. Notifications are stored in a '
    'dedicated table and displayed to the user in the Notifications page.'
)

# --- Non-Functional Requirements ---
doc.add_heading('3.3 Non-Functional Requirements', level=2)

doc.add_heading('Performance', level=3)
doc.add_paragraph(
    'The system should load the job listing page within three seconds under normal conditions. '
    'AI matching processes applicants one at a time with a 700ms delay between each so the admin can see '
    'real-time progress — this is intentional for demonstration purposes rather than performance optimization.'
)

doc.add_heading('Security', level=3)
sec_items = [
    'User passwords are hashed and handled entirely by Supabase Auth. The application never sees plaintext passwords.',
    'Row-Level Security (RLS) policies protect all database tables except "profiles" (RLS disabled for simplicity in this capstone).',
    'Sensitive operations (profile reads/writes, role checks) are wrapped in SECURITY DEFINER PostgreSQL functions so they bypass RLS with controlled access.',
    'The Gemini API key is stored in an environment variable (.env) and is never exposed to the client.',
]
for item in sec_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Reliability', level=3)
doc.add_paragraph(
    'The AI matching feature has a built-in fallback: if the Gemini API is unreachable or returns an error, '
    'the system switches to basic substring matching without crashing or showing errors to the user. '
    'The application is hosted on Supabase\'s infrastructure, which provides automatic backups and high availability.'
)

doc.add_heading('Usability', level=3)
usability_items = [
    'The login page uses a morph animation with an SVG checkmark and a personalized "Welcome back, [Name]!" message to make the sign-in experience feel responsive and polished.',
    'The AI analysis screen shows a progress bar with the current applicant\'s name so admins know the system is working.',
    'Cards use glass-morphism styling (semi-transparent backgrounds with backdrop blur) for a modern, clean look.',
    'Job cards in the applicant view have staggered entry animations for a polished feel.',
    'The search bar on Find Jobs includes a clear button to reset the search.',
]
for item in usability_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Maintainability', level=3)
doc.add_paragraph(
    'The codebase separates concerns into pages, components, context, and libraries. The AI matching logic lives '
    'in a single utility module (src/lib/aiSkillMatch.js) that can be swapped out or updated without touching '
    'any page code. The centralized theme file (uiStyles.js) means color and spacing changes propagate everywhere '
    'automatically.'
)

# Save
doc.save('C:\\Users\\admin\\LokalApp\\SRS_Lokal.docx')
print("SRS document generated: SRS_Lokal.docx")
